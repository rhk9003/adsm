import streamlit as st
import google.generativeai as genai
import tempfile
import os
import time
from pathlib import Path
from docx import Document
from io import BytesIO

# --- 頁面設定 ---
st.set_page_config(
    page_title="廣告策略 Gemini 3.0 生成器",
    page_icon="🎯",
    layout="wide",
    initial_sidebar_state="expanded"
)

# --- CSS 優化 ---
st.markdown("""
    <style>
    .stTextArea textarea { font-size: 14px; }
    .stButton button { width: 100%; border-radius: 8px; font-weight: bold; }
    .block-container { padding-top: 2rem; padding-bottom: 2rem; }
    div[data-testid="stExpander"] div[role="button"] p { font-size: 1.1rem; font-weight: 600; }
    </style>
""", unsafe_allow_html=True)

# --- Word 處理函式 ---

def extract_text_from_docx(file):
    """從上傳的 docx 檔案中提取文字"""
    try:
        doc = Document(file)
        full_text = []
        for para in doc.paragraphs:
            full_text.append(para.text)
        return '\n'.join(full_text)
    except Exception as e:
        st.error(f"Word 讀取失敗: {e}")
        return ""

def create_docx_from_markdown(markdown_text):
    """將 Markdown 文字轉換為 Word 檔案物件 (BytesIO)"""
    doc = Document()
    
    # 簡單的 Markdown 解析與寫入
    for line in markdown_text.split('\n'):
        line = line.strip()
        if not line:
            continue
            
        if line.startswith('# '):
            doc.add_heading(line[2:], level=1)
        elif line.startswith('## '):
            doc.add_heading(line[3:], level=2)
        elif line.startswith('### '):
            doc.add_heading(line[4:], level=3)
        elif line.startswith('- ') or line.startswith('* '):
            doc.add_paragraph(line[2:], style='List Bullet')
        else:
            # 移除簡單的 Markdown 標記以保持 Word 乾淨 (可選)
            clean_line = line.replace('**', '').replace('__', '')
            doc.add_paragraph(clean_line)
            
    # 儲存到記憶體中
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# --- Gemini 核心功能函式 ---

def configure_gemini(api_key):
    """設定 API Key"""
    if not api_key:
        st.error("❌ 請先在側邊欄輸入 Google Gemini API Key")
        return False
    try:
        genai.configure(api_key=api_key)
        return True
    except Exception as e:
        st.error(f"API Key 設定失敗: {e}")
        return False

def process_uploaded_file(uploaded_file):
    """
    處理上傳檔案：
    1. 若為 docx -> 轉為 txt 暫存檔 (因為 Gemini File API 原生不一定支援 docx 內容讀取，轉純文字最穩)
    2. 若為其他 -> 直接上傳
    回傳: Gemini File Object
    """
    if uploaded_file is None:
        return None
    
    try:
        suffix = Path(uploaded_file.name).suffix.lower()
        tmp_path = ""
        display_name = uploaded_file.name

        # 建立暫存檔案
        with tempfile.NamedTemporaryFile(delete=False, suffix=suffix) as tmp:
            tmp.write(uploaded_file.getvalue())
            tmp_path = tmp.name

        # 特殊處理 Word 檔：轉成 txt 再上傳給 Gemini，確保內容能被讀取
        if suffix == '.docx':
            text_content = extract_text_from_docx(tmp_path)
            os.remove(tmp_path) # 刪除原始 docx 暫存
            
            # 建立新的 txt 暫存檔
            with tempfile.NamedTemporaryFile(delete=False, suffix='.txt', mode='w', encoding='utf-8') as txt_tmp:
                txt_tmp.write(text_content)
                tmp_path = txt_tmp.name
                display_name = f"{uploaded_file.name}.txt"

        # 上傳至 Google GenAI
        with st.spinner(f"正在上傳並處理檔案: {uploaded_file.name} ..."):
            gemini_file = genai.upload_file(path=tmp_path, display_name=display_name)
            
            # 檢查處理狀態
            while gemini_file.state.name == "PROCESSING":
                time.sleep(1)
                gemini_file = genai.get_file(gemini_file.name)
            
            if gemini_file.state.name == "FAILED":
                st.error(f"檔案 {uploaded_file.name} 處理失敗。")
                return None
                
        # 清理本地暫存
        if os.path.exists(tmp_path):
            os.remove(tmp_path)
            
        return gemini_file

    except Exception as e:
        st.error(f"上傳錯誤 ({uploaded_file.name}): {e}")
        return None

def generate_content_stream(model_name, prompt, files=[]):
    """呼叫 Gemini API 生成內容"""
    try:
        model = genai.GenerativeModel(model_name)
        
        content_parts = [prompt]
        if files:
            content_parts.extend(files)
            
        with st.spinner(f"正在使用 {model_name} 模型進行深度運算中..."):
            response = model.generate_content(
                content_parts,
                generation_config=genai.types.GenerationConfig(
                    temperature=0.7,
                )
            )
        return response.text
    except Exception as e:
        st.error(f"生成錯誤: {e}")
        return None

# --- Session State 初始化 ---
if 'step1_result' not in st.session_state:
    st.session_state.step1_result = ""
if 'step2_result' not in st.session_state:
    st.session_state.step2_result = ""
if 'step3_result' not in st.session_state:
    st.session_state.step3_result = ""

# --- 側邊欄設定 ---
with st.sidebar:
    st.header("⚙️ 系統設定")
    api_key = st.text_input("輸入 Gemini API Key", type="password", help="請輸入您的 Google AI Studio API Key")
    
    st.markdown("### 🧠 模型選擇")
    model_options = [
        "gemini-3-pro",
        "gemini-3-pro-preview",
        "gemini-2.5-pro"
    ]
    selected_model = st.selectbox("使用模型", model_options, index=0)
    
    st.markdown("---")
    st.info(f"當前優先使用: **{selected_model}**")
    
    if st.button("重置所有分析", type="secondary"):
        st.session_state.step1_result = ""
        st.session_state.step2_result = ""
        st.session_state.step3_result = ""
        st.rerun()

# --- 主標題 ---
st.title("🎯 廣告策略 Gemini 3.0 智能工作台")
st.markdown("### 競品分析 → 差異比對 → 格式化素材產出 (Word 支援版)")
st.markdown("---")

# --- 分頁介面 ---
tab1, tab2, tab3 = st.tabs(["Step 1: 競品逆向工程", "Step 2: 我方現況比對", "Step 3: 格式化素材產出"])

# ==========================================
# Step 1: 競品深度分析
# ==========================================
with tab1:
    st.subheader("Step 1: 競品廣告庫分析")
    st.markdown("**目標**：上傳競爭對手廣告庫（PDF/Word/圖片/影片），產出戰略拆解報告。")
    
    col1, col2 = st.columns([1, 1])
    with col1:
        competitor_files = st.file_uploader(
            "上傳競品素材 (可多選: Word/PDF/圖/影)", 
            accept_multiple_files=True,
            type=['docx', 'png', 'jpg', 'jpeg', 'pdf', 'mp4', 'txt', 'csv'],
            key="s1_files"
        )
    with col2:
        competitor_text = st.text_area("直接貼上競品文案/連結 (選填)", height=150)

    if st.button("🚀 執行 Step 1 分析", type="primary", key="btn_s1"):
        if configure_gemini(api_key):
            gemini_files_s1 = []
            if competitor_files:
                for f in competitor_files:
                    g_file = process_uploaded_file(f)
                    if g_file: gemini_files_s1.append(g_file)
            
            prompt_s1 = f"""# Role: 資深廣告策略顧問
請針對我提供的【競爭對手廣告資料】(包含上傳的檔案與下方文字)進行深度逆向工程分析。

# 補充文字資料：
{competitor_text}

# 任務：產出《競品素材戰略拆解報告》
請嚴格依照以下 Markdown 架構分析：
1. **切角分類 (Hooks & Angles)**
2. **受眾推論 (Audience Profiling)**
3. **視覺與素材策略 (Visual Strategy)**
4. **文案語氣 (Tone & Manner)**
5. **素材套路庫 (Pattern Library)**
6. **我方戰略機會 (Strategic Gap)**

請給出詳盡、專業的分析報告。
"""
            result = generate_content_stream(selected_model, prompt_s1, gemini_files_s1)
            if result:
                st.session_state.step1_result = result
                st.success("Step 1 分析完成！")

    if st.session_state.step1_result:
        st.markdown("---")
        st.markdown("### 📝 Step 1 分析結果")
        st.markdown(st.session_state.step1_result)
        
        # Word 下載按鈕
        docx_file = create_docx_from_markdown(st.session_state.step1_result)
        st.download_button(
            label="📥 下載 Step 1 報告 (.docx)",
            data=docx_file,
            file_name="Step1_Competitor_Analysis.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )

# ==========================================
# Step 2: 我方現況比對
# ==========================================
with tab2:
    st.subheader("Step 2: 我方現有素材比對")
    
    if not st.session_state.step1_result:
        st.warning("⚠️ 請先完成 Step 1。")
    else:
        st.markdown("**目標**：基於 Step 1 的分析，檢視我方素材的缺口。")
        
        col1, col2 = st.columns([1, 1])
        with col1:
            our_files = st.file_uploader(
                "上傳我方現有素材 (可多選: Word/PDF/圖/影)", 
                accept_multiple_files=True,
                type=['docx', 'png', 'jpg', 'jpeg', 'pdf', 'mp4', 'txt'],
                key="s2_files"
            )
        with col2:
            our_text = st.text_area("補充我方資訊", height=150)

        if st.button("🚀 執行 Step 2 差異分析", type="primary", key="btn_s2"):
            if configure_gemini(api_key):
                gemini_files_s2 = []
                if our_files:
                    for f in our_files:
                        g_file = process_uploaded_file(f)
                        if g_file: gemini_files_s2.append(g_file)
                
                prompt_s2 = f"""# Context: 競品分析背景
{st.session_state.step1_result}

# Task: 差異化分析 (Gap Analysis)
請參考上述分析，並審視我現在上傳的【我方現有素材】(檔案) 以及下方補充資訊：
{our_text}

請進行比對並產出報告：
1. **現況盤點**
2. **盲區偵測 (The Gap)**
3. **優化建議**
4. **差異化突圍**
"""
                result = generate_content_stream(selected_model, prompt_s2, gemini_files_s2)
                if result:
                    st.session_state.step2_result = result
                    st.success("Step 2 比對完成！")

        if st.session_state.step2_result:
            st.markdown("---")
            st.markdown("### 📝 Step 2 分析結果")
            st.markdown(st.session_state.step2_result)
            
            # Word 下載按鈕
            docx_file_s2 = create_docx_from_markdown(st.session_state.step2_result)
            st.download_button(
                label="📥 下載 Step 2 報告 (.docx)",
                data=docx_file_s2,
                file_name="Step2_Gap_Analysis.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )

# ==========================================
# Step 3: 格式化素材產出
# ==========================================
with tab3:
    st.subheader("Step 3: 最終素材產出")
    
    if not st.session_state.step2_result:
        st.warning("⚠️ 請先完成 Step 1 與 Step 2。")
    else:
        st.markdown("**目標**：產出實際可用的文案與素材架構。")
        
        example_file = st.file_uploader(
            "上傳範例文件 (選填: Word/PDF)", 
            type=['docx', 'pdf', 'jpg', 'png', 'txt'],
            key="s3_example"
        )
        
        additional_req = st.text_input("額外要求", value="產出 4 組素材建議")

        if st.button("🚀 生成最終素材", type="primary", key="btn_s3"):
            if configure_gemini(api_key):
                gemini_files_s3 = []
                format_instruction = "請使用標準的文字條列格式。"
                
                if example_file:
                    g_file = process_uploaded_file(example_file)
                    if g_file:
                        gemini_files_s3.append(g_file)
                        format_instruction = "🚨 **格式嚴格要求**：請嚴格模仿附件檔案的「格式排版」與「欄位架構」。"
                
                prompt_s3 = f"""# Context
Step 1: {st.session_state.step1_result}
Step 2: {st.session_state.step2_result}

# Task: 創意素材產出
需求：{additional_req}。

# Format
{format_instruction}

# Requirements
1. **廣告主訴求**
2. **廣告素材文字**
3. **主文案**
4. **廣告標題**
"""
                result = generate_content_stream(selected_model, prompt_s3, gemini_files_s3)
                if result:
                    st.session_state.step3_result = result
                    st.success("Step 3 素材生成完成！")

        if st.session_state.step3_result:
            st.markdown("---")
            st.markdown("### 🎨 Step 3 素材建議")
            st.markdown(st.session_state.step3_result)
            
            # Word 下載按鈕
            docx_file_s3 = create_docx_from_markdown(st.session_state.step3_result)
            st.download_button(
                label="📥 下載最終素材檔案 (.docx)",
                data=docx_file_s3,
                file_name="Step3_Creative_Output.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )

# Footer
st.markdown("---")
st.caption(f"Powered by Google {selected_model} | Strategic Ad Toolkit")
